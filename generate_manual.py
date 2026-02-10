#!/usr/bin/env python3
"""
Script para gerar o Manual de Uso do Sistema 11Care - Plataforma de Enfermagem
Formato: DOCX
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ─── Page Setup ───────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

# ─── Styles ───────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Calibri'
    heading_style.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)
    if level == 1:
        heading_style.font.size = Pt(20)
        heading_style.paragraph_format.space_before = Pt(24)
        heading_style.paragraph_format.space_after = Pt(12)
    elif level == 2:
        heading_style.font.size = Pt(16)
        heading_style.paragraph_format.space_before = Pt(18)
        heading_style.paragraph_format.space_after = Pt(8)
    else:
        heading_style.font.size = Pt(13)
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)


def add_colored_paragraph(text, color=RGBColor(0x33, 0x33, 0x33), bold=False, size=Pt(11), align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = color
    run.bold = bold
    run.font.size = size
    if align:
        p.alignment = align
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * level)
    return p


def add_note_box(text, note_type="info"):
    colors = {
        "info": ("2196F3", "E3F2FD"),
        "warning": ("FF9800", "FFF3E0"),
        "danger": ("F44336", "FFEBEE"),
        "success": ("4CAF50", "E8F5E9"),
    }
    border_color, bg_color = colors.get(note_type, colors["info"])
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="24" w:space="8" w:color="{border_color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(borders)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}" w:val="clear"/>')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A568E" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shd)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            # Alternate row color
            if row_idx % 2 == 0:
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F8FC" w:val="clear"/>')
                cell._tc.get_or_add_tcPr().append(shd)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # spacing
    return table


# ═══════════════════════════════════════════════════════════════════════
# CAPA
# ═══════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('11Care')
run.bold = True
run.font.size = Pt(42)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Plataforma de Enfermagem')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x55, 0x88, 0xBB)
run.font.name = 'Calibri'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
# horizontal line
pPr = p._p.get_or_add_pPr()
borders = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="1A568E"/>'
    f'</w:pBdr>'
)
pPr.append(borders)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MANUAL DE USO DO SISTEMA')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Guia Completo para Profissionais de Enfermagem')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = 'Calibri'

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Versão 1.0')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Fevereiro de 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# Page break
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SUMARIO
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('Sumario', level=1)

toc_items = [
    ("1.", "Introducao", 3),
    ("1.1", "Sobre o Sistema 11Care", 3),
    ("1.2", "Publico-Alvo deste Manual", 3),
    ("1.3", "Requisitos do Sistema", 3),
    ("2.", "Acesso ao Sistema", 4),
    ("2.1", "Tela de Login", 4),
    ("2.2", "Primeiro Acesso e Troca de Senha", 4),
    ("2.3", "Recuperacao de Acesso", 5),
    ("3.", "Painel de Modulos", 5),
    ("3.1", "Visao Geral do Painel", 5),
    ("3.2", "Navegacao Principal", 5),
    ("4.", "Passagem de Plantao", 6),
    ("4.1", "Visao Geral", 6),
    ("4.2", "Lista de Pacientes", 6),
    ("4.3", "Busca e Filtros", 7),
    ("4.4", "Detalhes do Paciente", 7),
    ("4.5", "Metodologia SBAR", 8),
    ("4.6", "Anotacoes de Enfermagem", 8),
    ("4.7", "Classificacao de Risco", 9),
    ("4.8", "Analise Clinica com Inteligencia Artificial", 9),
    ("4.9", "Indicadores do Plantao", 10),
    ("4.10", "Impressao e Exportacao", 10),
    ("5.", "Sincronizacao de Dados", 11),
    ("5.1", "Sincronizacao Manual", 11),
    ("5.2", "Sincronizacao Automatica", 11),
    ("6.", "Historico de Pacientes", 11),
    ("6.1", "Consulta ao Historico", 11),
    ("6.2", "Reativacao de Pacientes", 12),
    ("7.", "Analytics e Relatorios", 12),
    ("7.1", "Dashboard de Analytics", 12),
    ("7.2", "Indicadores Disponiveis", 12),
    ("8.", "Painel Administrativo", 13),
    ("8.1", "Gerenciamento de Usuarios", 13),
    ("8.2", "Templates de Enfermarias", 14),
    ("8.3", "Gerenciamento de Enfermarias", 14),
    ("8.4", "Logs de Importacao", 15),
    ("8.5", "Analytics de Uso", 15),
    ("8.6", "Trilhas de Auditoria", 15),
    ("8.7", "Ferramentas de Manutencao", 16),
    ("9.", "Perfis de Acesso e Permissoes", 16),
    ("10.", "Boas Praticas e Recomendacoes", 17),
    ("11.", "Resolucao de Problemas Comuns", 17),
    ("12.", "Glossario", 18),
    ("13.", "Suporte Tecnico", 18),
]

for num, title, _page in toc_items:
    p = doc.add_paragraph()
    indent = 0
    if '.' in num and not num.endswith('.'):
        indent = 1
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(f'{num}  {title}')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    if num.endswith('.'):
        run.bold = True
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 1. INTRODUCAO
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('1. Introducao', level=1)

doc.add_heading('1.1 Sobre o Sistema 11Care', level=2)
doc.add_paragraph(
    'O 11Care e uma plataforma digital desenvolvida especificamente para profissionais '
    'de enfermagem, com foco na gestao de passagem de plantao e documentacao clinica. '
    'O sistema centraliza as informacoes dos pacientes, permitindo uma comunicacao '
    'eficiente e segura entre as equipes de enfermagem durante as trocas de turno.'
)
doc.add_paragraph(
    'A plataforma utiliza a metodologia SBAR (Situacao, Background/Historico, '
    'Avaliacao e Recomendacao), reconhecida internacionalmente como padrao de '
    'comunicacao clinica, garantindo que as informacoes essenciais sejam transmitidas '
    'de forma estruturada e completa.'
)
doc.add_paragraph('Principais funcionalidades do sistema:')
add_bullet('Gestao completa da passagem de plantao com metodologia SBAR')
add_bullet('Registro e acompanhamento detalhado de pacientes')
add_bullet('Classificacao de risco por semaforo (Vermelho, Amarelo, Verde)')
add_bullet('Analise clinica assistida por Inteligencia Artificial')
add_bullet('Indicadores e metricas do plantao em tempo real')
add_bullet('Historico completo de pacientes com arquivamento')
add_bullet('Exportacao de relatorios e impressao de passagem de plantao')
add_bullet('Painel administrativo com auditoria e controle de acesso')

doc.add_heading('1.2 Publico-Alvo deste Manual', level=2)
doc.add_paragraph(
    'Este manual destina-se a todos os profissionais de enfermagem que utilizarao '
    'o sistema 11Care no dia a dia, incluindo:'
)
add_bullet('Enfermeiros(as) assistenciais')
add_bullet('Tecnicos(as) de enfermagem')
add_bullet('Enfermeiros(as) lideres de equipe')
add_bullet('Coordenadores(as) de enfermagem')
add_bullet('Administradores do sistema')

doc.add_heading('1.3 Requisitos do Sistema', level=2)
doc.add_paragraph('Para utilizar o sistema 11Care, voce precisara de:')
add_table(
    ['Requisito', 'Especificacao'],
    [
        ['Navegador', 'Google Chrome (recomendado), Mozilla Firefox, Microsoft Edge - versoes atualizadas'],
        ['Conexao', 'Acesso a rede da instituicao (intranet ou internet)'],
        ['Dispositivo', 'Computador, notebook ou tablet'],
        ['Credenciais', 'Usuario e senha fornecidos pelo administrador do sistema'],
    ],
    col_widths=[4, 12]
)

add_note_box(
    'DICA: Recomendamos utilizar o navegador Google Chrome para a melhor '
    'experiencia de uso do sistema.',
    'info'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 2. ACESSO AO SISTEMA
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('2. Acesso ao Sistema', level=1)

doc.add_heading('2.1 Tela de Login', level=2)
doc.add_paragraph(
    'A tela de login e a primeira tela apresentada ao acessar o sistema 11Care. '
    'Siga os passos abaixo para realizar o acesso:'
)

doc.add_paragraph('Passo a passo para fazer login:', style='Normal')
p = doc.add_paragraph()
run = p.add_run('1. ')
run.bold = True
p.add_run('Abra o navegador e acesse o endereco (URL) do sistema fornecido pela sua instituicao.')
p = doc.add_paragraph()
run = p.add_run('2. ')
run.bold = True
p.add_run('Na tela de login, insira o seu nome de usuario no campo "Usuario".')
p = doc.add_paragraph()
run = p.add_run('3. ')
run.bold = True
p.add_run('Digite sua senha no campo "Senha".')
p = doc.add_paragraph()
run = p.add_run('4. ')
run.bold = True
p.add_run('Clique no botao "Entrar" para acessar o sistema.')

add_note_box(
    'IMPORTANTE: Suas credenciais sao pessoais e intransferiveis. Nunca compartilhe '
    'seu usuario e senha com outras pessoas. O sistema registra todas as acoes '
    'realizadas com sua conta para fins de auditoria.',
    'warning'
)

doc.add_heading('2.2 Primeiro Acesso e Troca de Senha', level=2)
doc.add_paragraph(
    'No primeiro acesso ao sistema, voce sera automaticamente direcionado(a) para a '
    'tela de troca de senha obrigatoria. Este procedimento e necessario para garantir '
    'a seguranca da sua conta.'
)

doc.add_paragraph('Requisitos para a nova senha:')
add_bullet('Minimo de 8 caracteres')
add_bullet('Deve conter letras e numeros')
add_bullet('Recomenda-se incluir letras maiusculas e minusculas')

doc.add_paragraph('Para alterar a senha:')
p = doc.add_paragraph()
run = p.add_run('1. ')
run.bold = True
p.add_run('Digite sua nova senha no campo indicado.')
p = doc.add_paragraph()
run = p.add_run('2. ')
run.bold = True
p.add_run('Confirme a nova senha digitando-a novamente.')
p = doc.add_paragraph()
run = p.add_run('3. ')
run.bold = True
p.add_run('Clique em "Alterar Senha" para confirmar.')

add_note_box(
    'ATENCAO: Apos a troca de senha, voce sera redirecionado(a) para a tela de login. '
    'Utilize a nova senha para acessar o sistema.',
    'warning'
)

doc.add_heading('2.3 Recuperacao de Acesso', level=2)
doc.add_paragraph(
    'Caso voce esqueca sua senha ou tenha dificuldades para acessar o sistema, '
    'entre em contato com o administrador do sistema da sua instituicao. '
    'O administrador podera redefinir sua senha, e voce precisara realizar '
    'a troca de senha novamente no proximo acesso.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 3. PAINEL DE MODULOS
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('3. Painel de Modulos', level=1)

doc.add_heading('3.1 Visao Geral do Painel', level=2)
doc.add_paragraph(
    'Apos o login, voce sera direcionado(a) para o Painel de Modulos. Esta e a tela '
    'principal do sistema, onde voce pode acessar todas as funcionalidades disponiveis '
    'de acordo com seu perfil de acesso.'
)
doc.add_paragraph('O painel apresenta os seguintes elementos:')
add_bullet('Cabecalho com o nome do sistema, mensagem de boas-vindas e botao de logout')
add_bullet('Cartoes de acesso rapido aos modulos do sistema')
add_bullet('Indicacao do seu perfil de acesso (Administrador, Enfermagem ou Visualizador)')

doc.add_heading('3.2 Navegacao Principal', level=2)
doc.add_paragraph(
    'A navegacao do sistema e organizada em modulos acessiveis a partir do painel principal:'
)

add_table(
    ['Modulo', 'Descricao', 'Acesso'],
    [
        ['Passagem de Plantao', 'Modulo principal para gestao da passagem de plantao e documentacao dos pacientes', 'Todos os perfis'],
        ['Painel Administrativo', 'Acesso as funcoes de administracao: usuarios, templates, enfermarias, auditoria', 'Apenas Administradores'],
    ],
    col_widths=[4.5, 8, 3.5]
)

doc.add_paragraph(
    'Para acessar um modulo, clique no cartao correspondente no painel. '
    'Para retornar ao painel de modulos a partir de qualquer tela, utilize '
    'o botao de retorno (seta) localizado no cabecalho da pagina.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 4. PASSAGEM DE PLANTAO
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('4. Passagem de Plantao', level=1)

doc.add_heading('4.1 Visao Geral', level=2)
doc.add_paragraph(
    'O modulo de Passagem de Plantao e o coracao do sistema 11Care. Ele permite que a equipe '
    'de enfermagem visualize, documente e transmita informacoes clinicas relevantes sobre '
    'cada paciente durante a troca de turno, seguindo a metodologia SBAR.'
)
doc.add_paragraph(
    'A tela principal do modulo apresenta:'
)
add_bullet('Barra de busca e filtros no topo da pagina')
add_bullet('Tabela com a lista de todos os pacientes internados')
add_bullet('Botoes de acao para sincronizacao, analise por IA, impressao e exportacao')
add_bullet('Indicadores resumidos do plantao')

doc.add_heading('4.2 Lista de Pacientes', level=2)
doc.add_paragraph(
    'A lista de pacientes apresenta todos os pacientes internados na enfermaria selecionada. '
    'Cada linha da tabela exibe informacoes resumidas do paciente, organizadas em colunas:'
)

add_table(
    ['Coluna', 'Informacao'],
    [
        ['Leito', 'Numero do leito em que o paciente esta internado'],
        ['Nome', 'Nome completo do paciente'],
        ['Registro', 'Numero de registro/prontuario do paciente'],
        ['Especialidade/Ramal', 'Especialidade medica e ramal de contato'],
        ['Diagnostico', 'Diagnostico principal do paciente'],
        ['Alerta', 'Indicadores visuais de alertas clinicos (coloridos por nivel de risco)'],
    ],
    col_widths=[4, 12]
)

doc.add_paragraph(
    'Clique em qualquer paciente na lista para abrir a tela de detalhes completos.'
)

doc.add_heading('4.3 Busca e Filtros', level=2)
doc.add_paragraph(
    'A barra de busca e filtros permite localizar pacientes rapidamente:'
)

p = doc.add_paragraph()
run = p.add_run('Busca por texto: ')
run.bold = True
p.add_run(
    'Digite no campo de busca o nome, numero do leito ou registro do paciente. '
    'A lista sera filtrada automaticamente conforme voce digita.'
)

p = doc.add_paragraph()
run = p.add_run('Filtro por alertas criticos: ')
run.bold = True
p.add_run(
    'Utilize o filtro de alertas para exibir apenas pacientes com alertas criticos '
    '(classificados em vermelho), facilitando a priorizacao do atendimento.'
)

p = doc.add_paragraph()
run = p.add_run('Filtro por enfermaria: ')
run.bold = True
p.add_run(
    'Selecione a enfermaria desejada no filtro para visualizar apenas os pacientes '
    'daquela unidade de internacao.'
)

doc.add_heading('4.4 Detalhes do Paciente', level=2)
doc.add_paragraph(
    'Ao clicar em um paciente na lista, um painel lateral (ou modal) sera aberto '
    'com todas as informacoes detalhadas. As informacoes sao organizadas conforme '
    'a metodologia SBAR e incluem:'
)

doc.add_heading('Dados de Identificacao', level=3)
add_bullet('Nome completo e registro do paciente')
add_bullet('Data de nascimento e idade')
add_bullet('Data de internacao e tempo de internamento')
add_bullet('Leito e enfermaria')
add_bullet('Especialidade medica e ramal')

doc.add_heading('Dados Clinicos', level=3)
add_bullet('Diagnostico principal')
add_bullet('Mobilidade (deambulando, acamado, restrito ao leito, etc.)')
add_bullet('Escala de Braden (avaliacao de risco para lesao por pressao)')
add_bullet('Alergias conhecidas')
add_bullet('Dieta prescrita')
add_bullet('Eliminacoes (diurese, evacuacao)')
add_bullet('Dispositivos invasivos (sonda vesical, acesso venoso central, etc.)')
add_bullet('Antibioticos em uso (ATB)')
add_bullet('Curativos e cuidados com feridas')
add_bullet('Aporte de oxigenio e saturacao')
add_bullet('Exames pendentes ou realizados')
add_bullet('Cirurgias programadas ou realizadas')
add_bullet('Previsao de alta')

doc.add_heading('Observacoes e Alertas', level=3)
add_bullet('Observacoes gerais de enfermagem')
add_bullet('Alertas clinicos categorizados por nivel de gravidade')

doc.add_heading('4.5 Metodologia SBAR', level=2)
doc.add_paragraph(
    'O sistema 11Care organiza as informacoes dos pacientes seguindo a metodologia SBAR, '
    'que e um padrao internacional de comunicacao clinica:'
)

add_table(
    ['Componente', 'Sigla', 'Descricao'],
    [
        ['Situacao', 'S', 'Identificacao do paciente, leito, diagnostico e motivo da comunicacao'],
        ['Background (Historico)', 'B', 'Historico clinico relevante, resultados de exames, medicamentos em uso'],
        ['Avaliacao', 'A', 'Estado atual do paciente, sinais vitais, alteracoes recentes, classificacao de risco'],
        ['Recomendacao', 'R', 'Plano de cuidados, encaminhamentos, exames pendentes, previsao de alta'],
    ],
    col_widths=[4, 1.5, 10.5]
)

add_note_box(
    'A metodologia SBAR e reconhecida pela Organizacao Mundial da Saude (OMS) e pela '
    'Joint Commission International como ferramenta essencial para a seguranca do paciente '
    'durante transicoes de cuidado.',
    'info'
)

doc.add_heading('4.6 Anotacoes de Enfermagem', level=2)
doc.add_paragraph(
    'O sistema permite que profissionais de enfermagem registrem anotacoes '
    'nao clinicas sobre cada paciente. Estas anotacoes servem para comunicar '
    'informacoes relevantes entre as equipes durante a troca de plantao.'
)

doc.add_paragraph('Para adicionar uma anotacao:')
p = doc.add_paragraph()
run = p.add_run('1. ')
run.bold = True
p.add_run('Abra os detalhes do paciente clicando em seu nome na lista.')
p = doc.add_paragraph()
run = p.add_run('2. ')
run.bold = True
p.add_run('Localize a secao "Anotacoes de Enfermagem".')
p = doc.add_paragraph()
run = p.add_run('3. ')
run.bold = True
p.add_run('Digite a anotacao no campo de texto (maximo de 200 caracteres).')
p = doc.add_paragraph()
run = p.add_run('4. ')
run.bold = True
p.add_run('Clique em "Salvar" para registrar a anotacao.')

doc.add_paragraph(
    'O sistema mantem um historico completo das anotacoes, incluindo o nome do '
    'profissional que registrou, a data e a hora. Anotacoes anteriores podem ser '
    'visualizadas no historico de anotacoes do paciente.'
)

add_note_box(
    'IMPORTANTE: Todas as anotacoes ficam registradas com identificacao do '
    'profissional responsavel, data e horario, para fins de rastreabilidade e auditoria.',
    'warning'
)

doc.add_heading('4.7 Classificacao de Risco', level=2)
doc.add_paragraph(
    'O sistema 11Care utiliza um sistema de classificacao por cores (semaforo) '
    'para indicar o nivel de atencao necessario para cada paciente:'
)

add_table(
    ['Cor', 'Nivel', 'Significado'],
    [
        ['VERMELHO', 'Critico', 'Paciente com condicao critica que requer atencao imediata. '
         'Pode indicar instabilidade clinica, dispositivos de alta complexidade ou risco elevado.'],
        ['AMARELO', 'Atencao', 'Paciente que necessita monitoramento continuo. '
         'Apresenta fatores de risco moderados ou pendencias importantes na documentacao.'],
        ['VERDE', 'Estavel', 'Paciente em condicao estavel, com documentacao completa '
         'e sem alertas criticos. Manter cuidados de rotina.'],
    ],
    col_widths=[3, 2.5, 10.5]
)

doc.add_paragraph(
    'A classificacao e exibida visualmente na lista de pacientes por meio de '
    'indicadores coloridos, permitindo uma rapida triagem visual durante a '
    'passagem de plantao.'
)

doc.add_heading('4.8 Analise Clinica com Inteligencia Artificial', level=2)
doc.add_paragraph(
    'O sistema conta com recurso de Inteligencia Artificial (IA) integrado, '
    'que oferece suporte a decisao clinica para a equipe de enfermagem. A IA analisa '
    'os dados do paciente e gera insights clinicos, incluindo:'
)

add_bullet('Score de qualidade da documentacao (0 a 100)')
add_bullet('Identificacao de lacunas criticas na documentacao')
add_bullet('Recomendacoes de cuidados de enfermagem')
add_bullet('Classificacao automatica do nivel de risco')
add_bullet('Sugestoes de prioridade de atendimento')

doc.add_paragraph('Para utilizar a analise por IA:')
p = doc.add_paragraph()
run = p.add_run('Analise Individual: ')
run.bold = True
p.add_run('Abra os detalhes do paciente e clique no botao "Analisar com IA". '
          'O sistema processara os dados e exibira os insights clinicos.')
p = doc.add_paragraph()
run = p.add_run('Analise em Lote: ')
run.bold = True
p.add_run('Na tela principal do plantao, clique no botao "Analise IA" para analisar '
          'todos os pacientes de uma vez. O sistema processara cada paciente e '
          'atualizara as classificacoes automaticamente.')

add_note_box(
    'ATENCAO: Os insights gerados pela IA sao ferramentas de apoio a decisao e NAO '
    'substituem o julgamento clinico do profissional de enfermagem. Sempre utilize '
    'sua avaliacao profissional em conjunto com as recomendacoes do sistema.',
    'danger'
)

doc.add_page_break()

doc.add_heading('4.9 Indicadores do Plantao', level=2)
doc.add_paragraph(
    'O sistema oferece uma visao consolidada dos indicadores do plantao atual, '
    'acessivel pela opcao "Visao Geral do Plantao" ou diretamente na tela de passagem. '
    'Os indicadores incluem:'
)

add_table(
    ['Indicador', 'Descricao'],
    [
        ['Total de Pacientes', 'Numero total de pacientes internados na enfermaria'],
        ['Media Braden', 'Media da Escala de Braden dos pacientes (risco de lesao por pressao)'],
        ['Media de Internacao', 'Tempo medio de internacao dos pacientes (em dias)'],
        ['Completude da Documentacao', 'Percentual de campos preenchidos nos registros dos pacientes'],
        ['Pacientes de Alta Complexidade', 'Numero de pacientes classificados como alto risco'],
        ['Pacientes com Dispositivos', 'Numero de pacientes com dispositivos invasivos'],
        ['Pacientes com ATB', 'Numero de pacientes em uso de antibioticos'],
        ['Pacientes Acamados', 'Numero de pacientes com mobilidade restrita ao leito'],
        ['Risco de Queda', 'Numero de pacientes com risco de queda identificado'],
        ['Risco de Lesao por Pressao', 'Numero de pacientes com risco conforme Escala de Braden'],
    ],
    col_widths=[5.5, 10.5]
)

doc.add_paragraph(
    'Esses indicadores auxiliam o enfermeiro lider a ter uma visao rapida do '
    'cenario do plantao e a direcionar a equipe conforme as prioridades.'
)

doc.add_heading('4.10 Impressao e Exportacao', level=2)
doc.add_paragraph('O sistema oferece duas opcoes para documentacao fisica:')

p = doc.add_paragraph()
run = p.add_run('Impressao da Passagem de Plantao: ')
run.bold = True
p.add_run(
    'Clique no botao "Imprimir" na tela do plantao. O sistema gerara um '
    'relatorio formatado com todos os dados dos pacientes, pronto para impressao. '
    'Este relatorio segue o formato SBAR e inclui todos os campos preenchidos.'
)

p = doc.add_paragraph()
run = p.add_run('Exportacao para Excel: ')
run.bold = True
p.add_run(
    'Clique no botao "Exportar Excel" para baixar uma planilha com os dados '
    'de todos os pacientes. A planilha pode ser utilizada para relatorios, '
    'analises estatisticas ou arquivamento.'
)

add_note_box(
    'LEMBRETE: Os documentos impressos e exportados contem dados sensiveis de '
    'pacientes. Mantenha-os em local seguro e descarte-os adequadamente conforme '
    'as politicas de seguranca da informacao da sua instituicao.',
    'warning'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 5. SINCRONIZACAO DE DADOS
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('5. Sincronizacao de Dados', level=1)

doc.add_heading('5.1 Sincronizacao Manual', level=2)
doc.add_paragraph(
    'O sistema permite sincronizar os dados dos pacientes com o sistema hospitalar '
    'de forma manual. Existem duas opcoes:'
)

p = doc.add_paragraph()
run = p.add_run('Sincronizar paciente individual: ')
run.bold = True
p.add_run(
    'Na tela de detalhes do paciente, clique no botao de sincronizacao para atualizar '
    'os dados daquele paciente especifico com as informacoes mais recentes do sistema hospitalar.'
)

p = doc.add_paragraph()
run = p.add_run('Sincronizar todos os pacientes: ')
run.bold = True
p.add_run(
    'Na tela principal do plantao, clique no botao "Sincronizar" para atualizar '
    'os dados de todos os pacientes da enfermaria.'
)

doc.add_paragraph(
    'Durante a sincronizacao, o sistema exibe indicadores visuais de progresso. '
    'Aguarde a conclusao antes de realizar outras operacoes.'
)

doc.add_heading('5.2 Sincronizacao Automatica', level=2)
doc.add_paragraph(
    'O sistema realiza sincronizacoes automaticas a cada hora, garantindo que os '
    'dados estejam sempre atualizados. Nao e necessario realizar a sincronizacao '
    'manual com frequencia, exceto quando voce precisar de dados em tempo real.'
)

add_note_box(
    'DICA: Recomenda-se realizar uma sincronizacao manual no inicio de cada plantao '
    'para garantir que voce esta trabalhando com os dados mais atualizados possiveis.',
    'info'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 6. HISTORICO DE PACIENTES
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('6. Historico de Pacientes', level=1)

doc.add_heading('6.1 Consulta ao Historico', level=2)
doc.add_paragraph(
    'O modulo de Historico de Pacientes armazena os registros de pacientes '
    'que receberam alta, foram transferidos ou tiveram seus registros arquivados. '
    'Para acessar o historico:'
)

p = doc.add_paragraph()
run = p.add_run('1. ')
run.bold = True
p.add_run('No painel de modulos, acesse "Historico de Pacientes".')
p = doc.add_paragraph()
run = p.add_run('2. ')
run.bold = True
p.add_run('Utilize os filtros disponiveis para localizar o paciente desejado:')
add_bullet('Nome do paciente')
add_bullet('Numero de registro')
add_bullet('Leito que ocupava')
add_bullet('Codigo de atendimento')
add_bullet('Motivo do arquivamento (alta, transferencia, etc.)')
add_bullet('Enfermaria')
add_bullet('Periodo (data inicial e final)')

p = doc.add_paragraph()
run = p.add_run('3. ')
run.bold = True
p.add_run('Clique no paciente para visualizar os detalhes completos do registro arquivado.')

doc.add_heading('6.2 Reativacao de Pacientes', level=2)
doc.add_paragraph(
    'Caso um paciente arquivado precise ser reativado (por exemplo, em caso de '
    'reinternacao), o sistema permite restaurar o registro:'
)

p = doc.add_paragraph()
run = p.add_run('1. ')
run.bold = True
p.add_run('Localize o paciente no historico.')
p = doc.add_paragraph()
run = p.add_run('2. ')
run.bold = True
p.add_run('Clique no botao "Reativar" no registro do paciente.')
p = doc.add_paragraph()
run = p.add_run('3. ')
run.bold = True
p.add_run('O paciente sera restaurado a lista ativa de pacientes da enfermaria.')

add_note_box(
    'ATENCAO: A reativacao de pacientes deve ser utilizada com criterio. '
    'Confirme com a equipe medica antes de reativar um registro.',
    'warning'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 7. ANALYTICS E RELATORIOS
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('7. Analytics e Relatorios', level=1)

doc.add_heading('7.1 Dashboard de Analytics', level=2)
doc.add_paragraph(
    'O modulo de Analytics oferece uma visao estatistica e grafica das informacoes '
    'do plantao. O dashboard apresenta graficos interativos e metricas que auxiliam '
    'na tomada de decisao e no monitoramento da qualidade assistencial.'
)

doc.add_heading('7.2 Indicadores Disponiveis', level=2)
doc.add_paragraph('Os principais indicadores disponiveis no Analytics sao:')

doc.add_heading('Metricas de Pacientes', level=3)
add_bullet('Distribuicao por especialidade medica')
add_bullet('Distribuicao por tipo de mobilidade')
add_bullet('Completude dos campos de documentacao')
add_bullet('Distribuicao por classificacao de risco')
add_bullet('Total de pacientes por enfermaria')

doc.add_heading('Metricas de Qualidade', level=3)
add_bullet('Percentual de completude da documentacao por campo')
add_bullet('Score medio de qualidade clinica')
add_bullet('Campos com maior indice de nao preenchimento')
add_bullet('Evolucao temporal dos indicadores')

doc.add_paragraph(
    'Os graficos sao interativos: passe o cursor sobre os elementos para ver '
    'detalhes e utilize os filtros para segmentar as informacoes por enfermaria, '
    'periodo ou outros criterios.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 8. PAINEL ADMINISTRATIVO
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('8. Painel Administrativo', level=1)

add_note_box(
    'NOTA: As funcoes descritas nesta secao estao disponiveis apenas para usuarios '
    'com perfil de Administrador.',
    'info'
)

doc.add_heading('8.1 Gerenciamento de Usuarios', level=2)
doc.add_paragraph(
    'O modulo de gerenciamento de usuarios permite aos administradores controlar '
    'o acesso ao sistema. As funcoes disponiveis sao:'
)

doc.add_heading('Criar Novo Usuario', level=3)
p = doc.add_paragraph()
run = p.add_run('1. ')
run.bold = True
p.add_run('Acesse "Painel Administrativo" > "Gerenciamento de Usuarios".')
p = doc.add_paragraph()
run = p.add_run('2. ')
run.bold = True
p.add_run('Clique no botao "Novo Usuario".')
p = doc.add_paragraph()
run = p.add_run('3. ')
run.bold = True
p.add_run('Preencha os campos obrigatorios:')
add_bullet('Nome completo do profissional')
add_bullet('Nome de usuario (login)')
add_bullet('E-mail institucional')
add_bullet('Senha inicial (o usuario devera alterar no primeiro acesso)')
add_bullet('Perfil de acesso (Administrador ou Enfermagem)')
p = doc.add_paragraph()
run = p.add_run('4. ')
run.bold = True
p.add_run('Clique em "Salvar" para criar o usuario.')

doc.add_heading('Editar Usuario', level=3)
doc.add_paragraph(
    'Na lista de usuarios, clique no botao de edicao ao lado do usuario desejado. '
    'Voce podera alterar o nome, e-mail, perfil de acesso e status de ativacao.'
)

doc.add_heading('Desativar/Ativar Usuario', level=3)
doc.add_paragraph(
    'Para desativar um usuario (por exemplo, em caso de desligamento), '
    'edite o usuario e altere o status para "Inativo". O usuario nao '
    'conseguira mais acessar o sistema, mas seus registros e anotacoes '
    'serao preservados para fins de auditoria.'
)

doc.add_heading('8.2 Templates de Enfermarias', level=2)
doc.add_paragraph(
    'Os templates permitem personalizar quais campos sao exibidos e quais sao '
    'obrigatorios para cada enfermaria. Isso permite adaptar o sistema as '
    'necessidades especificas de cada unidade de internacao.'
)

doc.add_paragraph('Para configurar um template:')
p = doc.add_paragraph()
run = p.add_run('1. ')
run.bold = True
p.add_run('Acesse "Painel Administrativo" > "Templates de Enfermarias".')
p = doc.add_paragraph()
run = p.add_run('2. ')
run.bold = True
p.add_run('Selecione a enfermaria desejada ou crie um novo template.')
p = doc.add_paragraph()
run = p.add_run('3. ')
run.bold = True
p.add_run('Para cada campo disponivel, defina:')
add_bullet('Visibilidade (exibir ou ocultar o campo)')
add_bullet('Obrigatoriedade (campo obrigatorio ou opcional)')
add_bullet('Regras especiais (se aplicavel)')
p = doc.add_paragraph()
run = p.add_run('4. ')
run.bold = True
p.add_run('Salve o template. As alteracoes serao aplicadas imediatamente para a enfermaria selecionada.')

doc.add_paragraph('Os campos configuraveis incluem:')
fields_table_data = [
    ['Leito', 'Numero do leito do paciente'],
    ['Nome', 'Nome completo do paciente'],
    ['Registro', 'Numero de prontuario'],
    ['Especialidade/Ramal', 'Especialidade medica e telefone'],
    ['Data de Nascimento', 'Data de nascimento do paciente'],
    ['Data de Internacao', 'Data de admissao'],
    ['Escala de Braden', 'Avaliacao de risco para lesao por pressao'],
    ['Diagnostico', 'Diagnostico clinico'],
    ['Mobilidade', 'Nivel de mobilidade do paciente'],
    ['Alergias', 'Alergias conhecidas'],
    ['Dieta', 'Tipo de dieta prescrita'],
    ['Eliminacoes', 'Informacoes sobre eliminacoes fisiologicas'],
    ['Dispositivos', 'Dispositivos invasivos em uso'],
    ['ATB', 'Antibioticos em uso'],
    ['Curativos', 'Cuidados com feridas e curativos'],
    ['Aporte/Saturacao', 'Suporte de oxigenio e saturacao'],
    ['Exames', 'Exames pendentes ou resultados'],
    ['Cirurgia', 'Cirurgias realizadas ou programadas'],
    ['Observacoes', 'Observacoes gerais'],
    ['Previsao de Alta', 'Data estimada de alta'],
    ['Alerta', 'Alertas clinicos'],
]
add_table(
    ['Campo', 'Descricao'],
    fields_table_data,
    col_widths=[4.5, 11.5]
)

doc.add_heading('8.3 Gerenciamento de Enfermarias', level=2)
doc.add_paragraph(
    'Este modulo permite cadastrar e gerenciar as enfermarias (unidades de internacao) '
    'do sistema. As funcoes incluem:'
)
add_bullet('Cadastrar novas enfermarias com codigo, nome, localizacao e descricao')
add_bullet('Sincronizar enfermarias com o sistema hospitalar (N8N)')
add_bullet('Aprovar ou rejeitar alteracoes pendentes nas enfermarias')
add_bullet('Ativar ou desativar enfermarias')
add_bullet('Editar informacoes das enfermarias existentes')

doc.add_paragraph(
    'Quando a sincronizacao com o sistema hospitalar detecta alteracoes '
    '(novas enfermarias ou mudancas em existentes), estas ficam como "pendentes" '
    'e precisam ser aprovadas pelo administrador antes de entrarem em vigor.'
)

doc.add_page_break()

doc.add_heading('8.4 Logs de Importacao', level=2)
doc.add_paragraph(
    'O modulo de Logs de Importacao exibe o historico completo das importacoes de dados '
    'realizadas pelo sistema, incluindo:'
)
add_bullet('Data e hora de cada importacao')
add_bullet('Enfermaria de origem dos dados')
add_bullet('Quantidade de registros processados')
add_bullet('Taxa de sucesso da importacao')
add_bullet('Erros encontrados durante o processo')
add_bullet('Duracao de cada importacao')

doc.add_paragraph(
    'Utilize os filtros para localizar importacoes especificas por periodo, '
    'enfermaria ou status (sucesso/erro).'
)

doc.add_heading('8.5 Analytics de Uso', level=2)
doc.add_paragraph(
    'O modulo de Analytics de Uso fornece metricas sobre como o sistema esta '
    'sendo utilizado pela equipe. Os indicadores incluem:'
)
add_bullet('Numero de sessoes ativas')
add_bullet('Duracao media das sessoes')
add_bullet('Paginas mais acessadas')
add_bullet('Acoes realizadas por usuario')
add_bullet('Dispositivos e navegadores utilizados')
add_bullet('Estatisticas por usuario')

doc.add_paragraph(
    'Essas informacoes auxiliam o administrador a identificar necessidades de '
    'treinamento, verificar a adocao do sistema e monitorar o engajamento da equipe.'
)

doc.add_heading('8.6 Trilhas de Auditoria', level=2)
doc.add_paragraph(
    'O sistema registra automaticamente todas as acoes realizadas pelos usuarios '
    'em uma trilha de auditoria completa. Cada registro inclui:'
)

add_table(
    ['Informacao', 'Descricao'],
    [
        ['Usuario', 'Nome e perfil do usuario que realizou a acao'],
        ['Acao', 'Tipo de acao realizada (criar, editar, excluir, visualizar, etc.)'],
        ['Recurso', 'Qual recurso foi afetado (paciente, usuario, enfermaria, etc.)'],
        ['Data/Hora', 'Data e hora exata da acao (formato UTC)'],
        ['Endereco IP', 'Endereco IP de origem da acao'],
        ['Detalhes', 'Informacoes adicionais sobre a acao realizada'],
    ],
    col_widths=[3.5, 12.5]
)

doc.add_paragraph('Tipos de acoes registradas:')
add_bullet('LOGIN / LOGOUT - Acesso e saida do sistema')
add_bullet('CREATE / UPDATE / DELETE - Criacao, alteracao e exclusao de registros')
add_bullet('EXPORT / IMPORT - Exportacao e importacao de dados')
add_bullet('SYNC_STARTED / SYNC_COMPLETED - Inicio e conclusao de sincronizacoes')
add_bullet('SHIFT_HANDOVER_VIEW / SHIFT_HANDOVER_PRINT - Visualizacao e impressao de plantao')
add_bullet('PATIENT_ARCHIVED / PATIENT_REACTIVATED - Arquivamento e reativacao de pacientes')

doc.add_heading('8.7 Ferramentas de Manutencao', level=2)
doc.add_paragraph(
    'O painel administrativo oferece ferramentas de manutencao do sistema:'
)

p = doc.add_paragraph()
run = p.add_run('Deduplicacao de Pacientes: ')
run.bold = True
p.add_run(
    'Remove registros duplicados de pacientes, mantendo o registro mais recente. '
    'Utilize esta ferramenta periodicamente para manter a base de dados limpa.'
)

p = doc.add_paragraph()
run = p.add_run('Limpeza de Orfaos: ')
run.bold = True
p.add_run(
    'Remove registros de pacientes que nao estao mais no sistema hospitalar '
    '(pacientes que receberam alta ou foram transferidos e ja nao constam no N8N).'
)

p = doc.add_paragraph()
run = p.add_run('Limpeza de Logs: ')
run.bold = True
p.add_run(
    'Permite remover logs de importacao antigos para manter o desempenho do sistema. '
    'Voce pode configurar o periodo de retencao em dias.'
)

add_note_box(
    'CUIDADO: As ferramentas de manutencao realizam operacoes irreversiveis. '
    'Execute-as com cautela e, preferencialmente, em horarios de baixo uso do sistema.',
    'danger'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 9. PERFIS DE ACESSO
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('9. Perfis de Acesso e Permissoes', level=1)

doc.add_paragraph(
    'O sistema 11Care possui tres perfis de acesso, cada um com permissoes '
    'especificas:'
)

add_table(
    ['Funcionalidade', 'Administrador', 'Enfermagem', 'Visualizador'],
    [
        ['Visualizar lista de pacientes', 'Sim', 'Sim', 'Sim'],
        ['Acessar detalhes do paciente', 'Sim', 'Sim', 'Sim'],
        ['Editar anotacoes de enfermagem', 'Sim', 'Sim', 'Nao'],
        ['Excluir anotacoes', 'Sim', 'Nao', 'Nao'],
        ['Executar analise por IA', 'Sim', 'Sim', 'Nao'],
        ['Imprimir passagem de plantao', 'Sim', 'Sim', 'Sim'],
        ['Exportar dados para Excel', 'Sim', 'Sim', 'Sim'],
        ['Acessar Analytics', 'Sim', 'Sim', 'Sim'],
        ['Acessar Historico de Pacientes', 'Sim', 'Sim', 'Sim'],
        ['Reativar pacientes arquivados', 'Sim', 'Sim', 'Nao'],
        ['Gerenciar usuarios', 'Sim', 'Nao', 'Nao'],
        ['Configurar templates', 'Sim', 'Nao', 'Nao'],
        ['Gerenciar enfermarias', 'Sim', 'Nao', 'Nao'],
        ['Acessar logs de importacao', 'Sim', 'Nao', 'Nao'],
        ['Visualizar trilhas de auditoria', 'Sim', 'Nao', 'Nao'],
        ['Acessar Analytics de Uso', 'Sim', 'Nao', 'Nao'],
        ['Executar ferramentas de manutencao', 'Sim', 'Nao', 'Nao'],
        ['Sincronizar dados com sistema hospitalar', 'Sim', 'Sim', 'Nao'],
    ],
    col_widths=[6, 3.3, 3.3, 3.3]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 10. BOAS PRATICAS
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('10. Boas Praticas e Recomendacoes', level=1)

doc.add_paragraph(
    'Para garantir o uso eficiente e seguro do sistema 11Care, recomendamos '
    'as seguintes boas praticas:'
)

doc.add_heading('Seguranca da Informacao', level=3)
add_bullet('Nunca compartilhe suas credenciais de acesso (usuario e senha) com terceiros')
add_bullet('Sempre faca logout ao encerrar o uso do sistema')
add_bullet('Nao deixe o sistema aberto em computadores de uso compartilhado sem supervisao')
add_bullet('Reporte imediatamente qualquer acesso nao autorizado ao administrador')

doc.add_heading('Qualidade da Documentacao', level=3)
add_bullet('Preencha todos os campos disponiveis no registro do paciente')
add_bullet('Utilize linguagem tecnica adequada e objetiva nas anotacoes')
add_bullet('Revise as informacoes antes de confirmar as alteracoes')
add_bullet('Registre anotacoes relevantes para a continuidade do cuidado')

doc.add_heading('Passagem de Plantao', level=3)
add_bullet('Realize uma sincronizacao manual no inicio de cada plantao')
add_bullet('Siga a estrutura SBAR para garantir uma comunicacao completa')
add_bullet('Priorize a comunicacao dos pacientes classificados em vermelho')
add_bullet('Utilize a analise por IA como ferramenta complementar')
add_bullet('Imprima ou exporte o relatorio do plantao como backup fisico')

doc.add_heading('Manutencao dos Dados (Administradores)', level=3)
add_bullet('Execute a deduplicacao de pacientes periodicamente')
add_bullet('Realize a limpeza de orfaos apos conferir as altas com o sistema hospitalar')
add_bullet('Monitore os logs de importacao para identificar falhas de sincronizacao')
add_bullet('Revise as trilhas de auditoria regularmente')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 11. RESOLUCAO DE PROBLEMAS
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('11. Resolucao de Problemas Comuns', level=1)

add_table(
    ['Problema', 'Causa Provavel', 'Solucao'],
    [
        ['Nao consigo fazer login',
         'Credenciais incorretas ou usuario desativado',
         'Verifique se o usuario e a senha estao corretos. Caso persista, contate o administrador para verificar o status da sua conta.'],
        ['A pagina nao carrega corretamente',
         'Problema de conexao ou cache do navegador',
         'Verifique sua conexao com a rede. Tente limpar o cache do navegador (Ctrl+Shift+Delete) e recarregar a pagina (F5).'],
        ['Os dados do paciente estao desatualizados',
         'Sincronizacao pendente',
         'Clique no botao "Sincronizar" para atualizar os dados com o sistema hospitalar.'],
        ['A analise por IA nao esta disponivel',
         'Servico de IA temporariamente indisponivel',
         'Aguarde alguns minutos e tente novamente. Caso persista, informe o administrador.'],
        ['Nao encontro um paciente na lista',
         'Paciente pode ter sido arquivado ou transferido',
         'Verifique o modulo "Historico de Pacientes". Se necessario, reative o paciente.'],
        ['A impressao nao funciona',
         'Configuracao do navegador ou impressora',
         'Verifique se o bloqueador de pop-ups do navegador esta desativado e se a impressora esta configurada corretamente.'],
        ['Recebi mensagem de sessao expirada',
         'Tempo de inatividade excedido',
         'Faca login novamente. O sistema encerra sessoes inativas automaticamente por seguranca.'],
        ['Nao consigo editar anotacoes',
         'Perfil sem permissao de edicao',
         'Verifique se seu perfil de acesso permite edicao. Perfis "Visualizador" nao possuem permissao de edicao.'],
    ],
    col_widths=[4, 4, 8]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 12. GLOSSARIO
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('12. Glossario', level=1)

add_table(
    ['Termo', 'Definicao'],
    [
        ['ATB', 'Antibioticoterapia - tratamento com antibioticos'],
        ['Braden (Escala de)', 'Escala utilizada para avaliar o risco de desenvolvimento de lesao por pressao. Quanto menor o score, maior o risco.'],
        ['Dashboard', 'Painel de controle visual que apresenta informacoes e indicadores de forma resumida e grafica'],
        ['Deduplicacao', 'Processo de remocao de registros duplicados no sistema, mantendo apenas o mais recente'],
        ['Dispositivos', 'Equipamentos ou materiais invasivos utilizados no paciente (sonda, cateter, dreno, etc.)'],
        ['Enfermaria', 'Unidade de internacao hospitalar onde os pacientes ficam alocados'],
        ['Evolucao', 'Registro da evolucao clinica do paciente, incluindo alteracoes no quadro e condutas realizadas'],
        ['IA (Inteligencia Artificial)', 'Tecnologia que analisa dados automaticamente e fornece insights e recomendacoes clinicas'],
        ['Leito', 'Local fisico (cama) onde o paciente esta internado, identificado por um numero'],
        ['N8N', 'Sistema de integracao utilizado para conectar o 11Care ao sistema hospitalar'],
        ['Orfao (paciente)', 'Registro de paciente que permanece no sistema mas nao consta mais no sistema hospitalar'],
        ['Passagem de Plantao', 'Processo de transferencia de informacoes entre equipes de enfermagem na troca de turno'],
        ['SBAR', 'Metodologia de comunicacao clinica: Situacao, Background, Avaliacao, Recomendacao'],
        ['Score', 'Pontuacao numerica utilizada para classificar ou avaliar uma condicao'],
        ['Sincronizacao', 'Processo de atualizacao dos dados entre o 11Care e o sistema hospitalar'],
        ['Template', 'Modelo de configuracao que define quais campos sao exibidos para cada enfermaria'],
        ['Trilha de Auditoria', 'Registro cronologico de todas as acoes realizadas no sistema para fins de rastreabilidade'],
    ],
    col_widths=[4.5, 11.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 13. SUPORTE TECNICO
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('13. Suporte Tecnico', level=1)

doc.add_paragraph(
    'Em caso de duvidas, dificuldades tecnicas ou sugestoes de melhoria, '
    'entre em contato com o suporte tecnico do sistema:'
)

doc.add_paragraph('Canais de atendimento:')
add_bullet('Administrador do sistema da sua instituicao (primeiro ponto de contato)')
add_bullet('Equipe de TI da instituicao (para problemas de infraestrutura e rede)')

doc.add_paragraph(
    'Ao solicitar suporte, informe:'
)
add_bullet('Seu nome e perfil de acesso')
add_bullet('Descricao detalhada do problema')
add_bullet('Tela ou funcionalidade em que o problema ocorre')
add_bullet('Mensagem de erro exibida (se houver)')
add_bullet('Navegador e dispositivo utilizados')

doc.add_paragraph()
doc.add_paragraph()

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p._p.get_or_add_pPr()
borders = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:top w:val="single" w:sz="6" w:space="8" w:color="CCCCCC"/>'
    f'</w:pBdr>'
)
pPr.append(borders)
run = p.add_run('11Care - Plataforma de Enfermagem')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Manual de Uso do Sistema - Versao 1.0 - Fevereiro de 2026')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Documento de uso interno - Distribuicao restrita a profissionais autorizados')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
run.font.name = 'Calibri'
run.italic = True

# ─── Save ─────────────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__), 'docs', 'Manual_de_Uso_11Care.docx')
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'Manual gerado com sucesso: {output_path}')
